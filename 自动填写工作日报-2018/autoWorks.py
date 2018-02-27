# python3
# autoWorks.py 自动填写周报。程序首先打开一个模板周报docx,然后打开一个现有的周报docx（名字格式为'工作日报-封晔-2018-2-27.docx'.名字里必须要有：如此格式的日期'2018-2-27'）
#              假如没有像如此格式的'工作日报-封晔-2018-2-27.docx'文档，必须先手动创建一个，里面还要厚模板周报里的表格

import docx
import os
import datetime
import re
import subprocess 
from docx.enum.text import WD_ALIGN_PARAGRAPH

#模板工作日志docx位置
templateWorkLog = os.path.join(os.getcwd(), '1工作日报-模板.docx')
#当天日期%Y-%m-%d
dateTimeNow = datetime.datetime.now().strftime('%Y-%m-%d')

def writeTemplateWorkDocx():
    '''
    把今日工作内容填写到模板工作日志里
    '''
    doc = docx.Document(templateWorkLog) #读取模板docx文件
    tabs = doc.tables[0] #读取模板表格
    tabs.rows[0].cells[1].text = dateTimeNow #工作日报日期时间
    print('请输入任务内容：')
    inputTask = input()
    tabs.rows[1].cells[1].text = inputTask #输入今日任务
    print('请输入完成情况')
    inputPercent = input()
    tabs.rows[2].cells[1].text = inputPercent #输入今日完成情况
    print('请输入明日工作')
    tomorrowTask = input()
    tabs.rows[3].cells[1].text = tomorrowTask #输入明日工作
    return doc

def getLatestDateWorkLogDocx():
    '''
    获取日期最近的已经填写的日志文件
    '''
    doc = writeTemplateWorkDocx() #获取模板docx
    latestDoc = docx.Document(os.path.join(os.getcwd(), getMaxDateTimeDocx())) #获取已填写的工作日志的doc
    
    for tb in latestDoc.tables: #遍历表
        rowsCount = len(tb.rows) #当前表格的行数
        colCount = len(tb.columns)   #当前表格的列数
        newTable = doc.add_table(rowsCount, colCount, tb.style) #创建新表格
        newTable.alignment = WD_ALIGN_PARAGRAPH.CENTER #设置表格居中
        for tbRow in range(rowsCount): #遍历表中的行
            for tbCol in range(colCount): #遍历表中的列
                newTable.cell(tbRow,tbCol).width = tb.cell(tbRow,tbCol).width #设置表格列宽
                newTable.rows[tbRow].cells[tbCol].text = tb.rows[tbRow].cells[tbCol].text #设置单元格内容
    
    return doc

def getMaxDateTimeDocx():
    '''
    获取日期最大的docx文件
    '''
    L = []
    fileName = ''
    for f in os.listdir('.'):
        if f.endswith('docx') or f.endswith('doc'):
            m = re.search(r'\d{4}-\d{2}-\d{2}', f) #要与dateTimeNow格式化的日期一致
            if m is not None:
                L.append(m.group())
    L.sort()
    str1 = L[len(L) - 1]
    for f in os.listdir('.'):
        if str(str1) in str(f):
            fileName = f
            break
    return fileName

def startAutoWorks():
    '''
    启动程序
    '''
    try:
        docResult = getLatestDateWorkLogDocx()
        saveFile = os.path.join(os.getcwd(), '工作日报-封晔-' + dateTimeNow + '.docx')
        docResult.save(saveFile)
        subprocess.Popen(['start', saveFile], shell=True) #打开完档        return 'ok'
    except Exception as e:
        return e

#执行程序 startAutoWorks()
print(startAutoWorks())